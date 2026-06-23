import streamlit as st
from datetime import datetime
from io import BytesIO

from docx import Document
from reportlab.pdfgen import canvas


def generar_word(data):

    doc = Document()

    doc.add_heading(
        "INFORME DE EVALUACIÓN CREDITICIA",
        level=1
    )

    for seccion, contenido in data.items():

        doc.add_heading(
            seccion,
            level=2
        )

        for k, v in contenido.items():

            doc.add_paragraph(
                f"{k}: {v}"
            )

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer


def generar_pdf(data):

    buffer = BytesIO()

    pdf = canvas.Canvas(buffer)

    y = 800

    pdf.setFont(
        "Helvetica-Bold",
        14
    )

    pdf.drawString(
        50,
        y,
        "INFORME DE EVALUACIÓN CREDITICIA"
    )

    y -= 30

    pdf.setFont(
        "Helvetica",
        9
    )

    for seccion, contenido in data.items():

        pdf.drawString(
            50,
            y,
            f"{seccion}"
        )

        y -= 20

        for k, v in contenido.items():

            pdf.drawString(
                60,
                y,
                f"{k}: {v}"
            )

            y -= 15

            if y < 50:

                pdf.showPage()

                y = 800

    pdf.save()

    buffer.seek(0)

    return buffer


def render():

    st.subheader("📄 Reporte Final")

    cliente = st.session_state.get(
        "cliente_actual",
        {}
    )

    evaluacion = st.session_state.get(
        "evaluacion_credito",
        {}
    )

    ficha = st.session_state.get(
        "ficha_cliente",
        {}
    )

    finanzas = st.session_state.get(
        "ingresos_gastos",
        {}
    )

    visita = st.session_state.get(
        "ubicacion_visita",
        {}
    )

    usuario = st.session_state.get(
        "usuario",
        "SIN_USUARIO"
    )

    # -------------------------------
    # Consolidado
    # -------------------------------

    reporte = {

        "CLIENTE": cliente,

        "EVALUACION": evaluacion,

        "FICHA": ficha,

        "FINANZAS": finanzas,

        "VISITA": visita
    }

    # -------------------------------
    # Vista previa
    # -------------------------------

    st.markdown(
        "## Vista Previa"
    )

    with st.expander(
        "Ver Información Consolidada"
    ):

        st.json(reporte)

    st.divider()

    # -------------------------------
    # Resumen ejecutivo
    # -------------------------------

    st.markdown(
        "## Resumen"
    )

    st.write(
        f"Cliente: {cliente.get('CLIENTE','')}"
    )

    st.write(
        f"DNI: {cliente.get('DOCPEN','')}"
    )

    st.write(
        f"Crédito: {cliente.get('CODCRE','')}"
    )

    st.write(
        f"Riesgo: {evaluacion.get('riesgo','')}"
    )

    st.write(
        f"Resultado Crédito: "
        f"{evaluacion.get('recomendacion','')}"
    )

    st.write(
        f"Resultado Financiero: "
        f"{finanzas.get('conclusion_financiera','')}"
    )

    st.write(
        f"Resultado Visita: "
        f"{visita.get('resultado_visita','')}"
    )

    st.divider()

    # -------------------------------
    # DOCX
    # -------------------------------

    if st.button(
        "📄 Generar Word"
    ):

        archivo_word = generar_word(
            reporte
        )

        st.download_button(
            label="⬇ Descargar DOCX",
            data=archivo_word,
            file_name=
            f"Informe_{cliente.get('CODCRE','')}.docx",
            mime=
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # -------------------------------
    # PDF
    # -------------------------------

    if st.button(
        "📑 Generar PDF"
    ):

        archivo_pdf = generar_pdf(
            reporte
        )

        st.download_button(
            label="⬇ Descargar PDF",
            data=archivo_pdf,
            file_name=
            f"Informe_{cliente.get('CODCRE','')}.pdf",
            mime="application/pdf"
        )

    st.divider()

    # -------------------------------
    # Historial
    # -------------------------------

    registro = {

        "usuario": usuario,

        "fecha":
        datetime.now().strftime(
            "%Y-%m-%d %H:%M:%S"
        ),

        "cliente":
        cliente.get(
            "CLIENTE",
            ""
        ),

        "dni":
        cliente.get(
            "DOCPEN",
            ""
        ),

        "credito":
        cliente.get(
            "CODCRE",
            ""
        ),

        "resultado_credito":
        evaluacion.get(
            "recomendacion",
            ""
        ),

        "resultado_financiero":
        finanzas.get(
            "conclusion_financiera",
            ""
        ),

        "resultado_visita":
        visita.get(
            "resultado_visita",
            ""
        )
    }

    st.markdown(
        "## Registro"
    )

    st.json(registro)

    st.success(
        "Reporte listo para exportar."
    )
